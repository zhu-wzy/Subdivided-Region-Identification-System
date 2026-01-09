from matplotlib import pyplot as plt
from sklearn import metrics
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import GridSearchCV, cross_val_score, train_test_split
from sklearn.metrics import mean_squared_error,explained_variance_score,mean_absolute_error, roc_auc_score, roc_curve
import pydotplus
from sklearn.tree import export_graphviz
import joblib

from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA

import numpy as np
import math
import time
import pandas as pd
from pylab import mpl
import os

from docx import Document
from docx.oxml.ns import qn


##配置参数
 
#模式类
pkl_traim_select      = '高速非高速'       # 2个训练模型选择        '高速非高速'        '郊区山区'
pkl_traim_mode_select = '简单调参局部最优'  # 调参的方法            '简单调参局部最优'  '网格搜索全局最优'
pkl_traim_test_select = '模型训练模式'     #  模型训练或者模型验证  '模型训练模式'      '模型验证模式'
 
#取值范围类  简单调参局部最优局部最优取值范围
n_estimators_range      = [0,200] #决策树个数
n_estimators_range_step = 10      #大网格细分间隔

max_depth_range         = [1, 20] #决策树深度
max_depth_range_step    = 1       #大网格细分间隔

#max_features_range      = [特征参数开根号,特征参数]    #最大特征值
max_features_range_step = 1       #大网格细分间隔

min_samples_leaf_range  = 1 #叶节点必须有的最小样本数量
min_samples_split_range = 2 #分裂前节点必须有的最小样本数

#网格搜索全局最优 全局整体寻优索引范围
global_range=3   #局部寻优后全部变量一起在最优值附近寻优前后索引范围    3:  最优值-3~最优值+3



if pkl_traim_select=='高速非高速':
    class_names=['高速','非高速']
    filepath = 'data_train_car_gaosu'  
if pkl_traim_select=='郊区山区':
    class_names=['郊区','山区']  
    filepath = 'data_train_car_jiaoqu'  

if pkl_traim_mode_select == '简单调参局部最优':  
    tiaocan_id = 0
if pkl_traim_mode_select == '网格搜索全局最优':   
    tiaocan_id = 1

if pkl_traim_test_select == '模型训练模式':
    xunlian_flag = 0
if pkl_traim_test_select == '模型验证模式':
    xunlian_flag = 1

tiaocanmoshi={0:'0-简单调参局部最优/',1:'1-网格搜索全局最优/'}

#添加输出路径
def mkdir(path):
    # 去除首位空格./ 
    path=path.strip(' ./')
    # 去除尾部 \ 符号
    path=path.rstrip("\\")
 
    # 判断路径是否存在
    # 存在     True
    # 不存在   False
    isExists=os.path.exists(path)

    # 判断结果
    if not isExists:
        # 如果不存在则创建目录
        # 创建目录操作函数
        os.makedirs(path) 
 
        #print (path+' 创建成功')
        return True
    else:
        # 如果目录存在则不创建，并提示目录已存在
        #print (path+' 目录已存在')
        return False

#保存模型
def save_model(model, filepath):
    joblib.dump(model, filename=filepath)
#读取模型
def load_model(filepath):
    model = joblib.load(filepath)
    return model

#导入数据
def data_divide(filepath,xunlian_flag):
    ###读取csv
    data=pd.read_csv('.\\'+filepath+'.csv',index_col=0,header=0,encoding='gb2312') 

    ###去除重复值
    print('存在' if any(data.duplicated()) else '不存在', '重复观测值')
    data=data.drop_duplicates()
    
    ###特征值
    feature_name =list(data.columns)[0:-1]
    max_features = len(feature_name)

    ###划分训练集、验证集
    if xunlian_flag == 0:
        data_train, data_test= train_test_split(data,test_size=0.3, random_state=0)
       
        print ("训练集统计描述：\n",data_train.describe().round(2))
        print ("验证集统计描述：\n",data_test.describe().round(2))
        print ("训练集信息：\n",data_train.iloc[:,-1].value_counts())
        print ("验证集信息：\n",data_test.iloc[:,-1].value_counts())

        x_train=data_train.iloc[:,0:-1].values.astype('float32') 
        x_test=data_test.iloc[:,0:-1].values.astype('float32')
        y_train=data_train.iloc[:,-1].values.astype('float32')
        y_test=data_test.iloc[:,-1].values.astype('float32')

        return x_train,x_test,y_train,y_test,data_train, data_test,feature_name,max_features

    else :
        x_test=data.iloc[:,0:-1].values.astype('float32')
        y_test=data.iloc[:,-1].values.astype('float32')
        return x_test,y_test,data,feature_name,max_features

#调参
def tiaocan(wanggesousuo): ###0简单调参 1网格调参
    star=time.perf_counter()

    ###大体规划决策树个数n_estimators
    scorel = []
    for i in range(n_estimators_range[0],n_estimators_range[1],n_estimators_range_step):
        rfc = RandomForestClassifier(n_estimators=i+1,n_jobs=-1,random_state=0)
        score = cross_val_score(rfc,x_train,y_train,cv=7).mean()
        scorel.append(score)
    roughly_n_estimators_id = (scorel.index(max(scorel))*n_estimators_range_step)+1
    print('大体规划n_estimators:')
    print('准确度：',max(scorel),' 结果：',roughly_n_estimators_id)

    ###确定决策树个数n_estimators
    scorel = []
    for i in range(roughly_n_estimators_id-n_estimators_range_step,roughly_n_estimators_id+n_estimators_range_step):
        rfc = RandomForestClassifier(n_estimators=i,n_jobs=-1,random_state=0)
        score = cross_val_score(rfc,x_train,y_train,cv=7).mean()
        scorel.append(score)
    n_estimators_id = ([*range(roughly_n_estimators_id-n_estimators_range_step,roughly_n_estimators_id+n_estimators_range_step)][scorel.index(max(scorel))])
    print('确定决策树个数n_estimators')
    print('准确度：',max(scorel),' 结果：',n_estimators_id) 




    ###确定决策树深度max_depth
    param_grid = {'max_depth':np.arange(max_depth_range[0], max_depth_range[1], max_depth_range_step)}
    rfc = RandomForestClassifier(n_estimators=n_estimators_id,n_jobs=-1,random_state=0)
    GS = GridSearchCV(rfc,param_grid,cv=7)
    GS.fit(x_train,y_train)
    max_depth_id = GS.best_params_['max_depth']
    print('确定决策树深度max_depth')
    print('准确度：',GS.best_score_,' 结果：',GS.best_params_)     




    ###确定最大特征值max_features
    param_grid = {'max_features':np.arange(int(math.sqrt(max_features)),max_features+1,max_features_range_step)}#特征参数开根号~特征参数  +1是range需要
    rfc = RandomForestClassifier(n_estimators=n_estimators_id,random_state=0,n_jobs=-1)
    GS = GridSearchCV(rfc,param_grid,cv=7)
    GS.fit(x_train,y_train)
    max_features_id = GS.best_params_['max_features']
    print('确定最大特征值max_features')
    print('准确度：',GS.best_score_,' 结果：',GS.best_params_)  
    end=time.perf_counter()

    print("局部优化训练时间：{:.2f}s  {:.2f}h".format(end-star,(end-star)/3600))
    print('局部优化搜索结果')
    print('准确度：',GS.best_score_)  
    
    ###写入Word
    try :
        document = Document('./result/Word/参数.docx')
    except :
        document = Document()
        document.save('./result/Word/参数.docx')
        document = Document('./result/Word/参数.docx')
        paragraph = document.add_paragraph('调参参数：\n')

    time_id = time.localtime()
    time_now = time.strftime("%Y-%m-%d %H:%M:%S", time_id)
    paragraph = document.add_paragraph(time_now)
    paragraph.add_run('\n{}\n'.format(filepath))
    paragraph.add_run('局部最优：\n')    
    paragraph.add_run('n_estimators:{} max_depth:{} max_features:{}\n '.format(n_estimators_id,max_depth_id,max_features_id))

    ###网格搜索
    if wanggesousuo == 1:
        star=time.perf_counter()
        print('开始网格搜索')
        try:
            n_estimators_id>3
            max_depth_id>3
            max_features_id>3
        except:
            print('参数过小，修改网格搜索参数范围,修改后请屏蔽try')
 
        param_grid={'n_estimators':np.arange(n_estimators_id-global_range,n_estimators_id+global_range, 1),
                    'max_depth':np.arange(max_depth_id-global_range,max_depth_id+global_range, 1),
                    'max_features':np.arange(max_features_id-global_range,max_features_id+global_range,1),
                    'min_samples_leaf':np.arange(min_samples_leaf_range,min_samples_leaf_range+global_range, 1),
                    'min_samples_split':np.arange(min_samples_split_range,min_samples_split_range+global_range, 1),
                    'criterion':["gini", "entropy"]
                    }
        rfc = RandomForestClassifier(random_state=0,n_jobs=-1)
        GS = GridSearchCV(rfc,param_grid,cv=10)
        GS.fit(x_train,y_train)
        end=time.perf_counter()
        print("网格搜索训练时间：{:.2f}s  {:.2f}h".format(end-star,(end-star)/3600))
        print('网格搜索结果')
        print('准确度：',GS.best_score_,' 结果：',GS.best_params_)  
        paragraph.add_run('网格搜索：{}\n'.format(bool(wanggesousuo)))    
        n_estimators_id=GS.best_params_['n_estimators']
        max_depth_id=GS.best_params_['max_depth']
        max_features_id=GS.best_params_['max_features']
        min_samples_leaf_id = GS.best_params_['min_samples_leaf']
        min_samples_split_id = GS.best_params_['min_samples_split']
        criterion_id = GS.best_params_['criterion']
        paragraph.add_run('n_estimators:{} max_depth:{} max_features:{} min_samples_leaf:{} min_samples_split:{} criterion:{}\n '
                         .format(n_estimators_id,max_depth_id,max_features_id,min_samples_leaf_id,min_samples_split_id,criterion_id))
    else :
        min_samples_split_id = 2
        min_samples_leaf_id = 1
        criterion_id = "gini"                   

    document.save('./result/Word/参数.docx')
    para = {'n_estimators':n_estimators_id,'max_depth':max_depth_id,'max_features':max_features_id,
           'min_samples_leaf':min_samples_leaf_id,'min_samples_split':min_samples_split_id,'criterion':criterion_id}
    return para

#输出标签
def biaoqian(y_pdt,probility):
    paragraph = document.add_paragraph('预测标签：\n')
    paragraph.add_run(str(list(y_pdt))+'\n')
    paragraph.add_run('测试集标签：\n')
    paragraph.add_run(str(list(y_test))+ '\n')
    paragraph.add_run('ID:\n')
    paragraph.add_run(str(list(data_test.index))+ '\n')  ###输出对应ID（应用模型时是全部ID）
    paragraph.add_run('分类概率：\n')
    paragraph.add_run(str(list(probility))+ '\n')

    document.save('./result/Word/标签.docx')

#训练模型
def model_create(n_estimators,max_depth,max_features,min_samples_split,min_samples_leaf,criterion):
    NBM = [RandomForestClassifier(n_estimators=n_estimators,max_depth=max_depth,max_features=max_features,random_state=0,
                                  min_samples_split=min_samples_split,min_samples_leaf=min_samples_leaf,criterion=criterion)]
    NAME = ["RF"]
    for itr, itrname in zip(NBM, NAME):
        print("Training...")
        itr.fit(x_train, y_train)
         
        ###保存模型
        filepath_temp = './result/pkl/'+tiaocanmoshi[tiaocan_id]+filepath+'.pkl'
        save_model(itr,filepath_temp)

#预测结果
def use_model(itrname,xunlian_flag):
        ###加载模型
        print("Applying...")

        if xunlian_flag == 0:
            filepath_temp = './result/pkl/'+tiaocanmoshi[tiaocan_id]+filepath+'.pkl'
        else:
            filepath_temp = './result/pkl/'+tiaocanmoshi[tiaocan_id]+pklpath+'.pkl'
        itr=load_model(filepath_temp)

        ###预测
        y_pdt = itr.predict(x_test)       
        ###输出标签
        probility = itr.predict_proba(x_test)
        biaoqian(y_pdt,probility)

        if xunlian_flag == 0 :
            y_train_pdt = itr.predict(x_train)
            
            #重要性
            importances = itr.feature_importances_
            print("重要性：",importances)
            indices = np.argsort(importances)[::-1]
            document = Document()
            paragraph = document.add_paragraph('重要性：\n')
            for f in range(x_train.shape[1]):
                paragraph.add_run("%2d) %-*s %f \n" % (f + 1, 30, feature_name[indices[f]], importances[indices[f]]))
                print("%2d) %-*s %f" % (f + 1, 30, feature_name[indices[f]], importances[indices[f]]))
            document.save('./result/Word/重要性.docx')
            ###输出结果
            dts1 = len(np.where(y_train_pdt==y_train)[0])/len(y_train)
            dts2 = len(np.where(y_pdt==y_test)[0])/len(y_test)  
           
            print("训练集：{}             ".format(itrname,))
            print("训练集：{} 精度:{:.3f}%".format(itrname, dts1*100))
            print("测试集：{} 精度:{:.3f}%".format(itrname, dts2*100))

            ###ROC曲线 AUC值
            auc = roc_auc_score(y_test,itr.predict_proba(x_test)[:,1])
            # auc = roc_auc_score(y_test,clf.decision_function(X_test))
            fpr,tpr, thresholds = roc_curve(y_test,itr.predict_proba(x_test)[:,1])
            plt.plot(fpr,tpr,color='darkorange',label='ROC curve (area = %0.3f)' % auc)
            plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
            plt.xlim([0.0, 1.0])
            plt.ylim([0.0, 1.05])
            plt.xlabel('False Positive Rate')
            plt.ylabel('True Positive Rate')
            plt.title('Receiver operating characteristic example')
            plt.legend(loc="lower right")
            plt.savefig('./result/图片/{0}_{1}_ROC曲线.jpg'.format(pklpath,tiaocan_id),dpi=800)      ###dpi 像素
    
        else :
            ###输出结果
            dts2 = len(np.where(y_pdt==y_test)[0])/len(y_test)  
            print("测试集：{} 精度:{:.3f}%".format(itrname, dts2*100))
            
        ###叶节点准确率
        right_1=right_2=0
        sum_1=len(np.extract(y_pdt==0,y_pdt))
        sum_2=len(np.extract(y_pdt==1,y_pdt))

        call_1=call_2=0
        sum_c1=len(np.extract(y_test==0,y_test)) 
        sum_c2=len(np.extract(y_test==1,y_test)) 

        xiangdeng=np.where(y_pdt==y_test)[0] 

        for i in xiangdeng :
            if y_test[i] ==0: 
                right_1+=1
                call_1+=1
            if y_test[i] ==1: 
                right_2+=1
                call_2+=1                          

        if sum_1!=0 :
            acc_1=right_1/sum_1*100
        else: 
            acc_1=0
        if sum_2!=0 :
            acc_2=right_2/sum_2*100
        else: 
            acc_2=0         

        Rec_1=call_1/sum_c1*100
        Rec_2=call_2/sum_c2*100
    
        print("测试集各类精确度: {0}: {1:.3f}%  {2}:{3:.3f}% ".format(class_names[0],acc_1,class_names[1],acc_2))
        print("测试集各类召回率: {0}: {1:.3f}%  {2}:{3:.3f}% ".format(class_names[0],Rec_1,class_names[1],Rec_2))

        disp = metrics.ConfusionMatrixDisplay.from_estimator(itr, x_test, y_test)    ###混淆矩阵  
        disp.figure_.suptitle("Confusion Matrix")
        print("Confusion matrix:\n%s" % disp.confusion_matrix)

#可视化
def keshihua(itr):
    ###模型可视化
    print('输出树')
    # 循环打印每棵树
    for idx, estimator in enumerate(itr.estimators_):
        # 导出dot文件
        dot_data=export_graphviz(estimator,
                        out_file=None,
                        feature_names=feature_name,
                        class_names=class_names,
                        rounded=True,
                        special_characters=True,
                        precision=4,
                        fontname="NSimSun",
                        proportion=False, 
                        filled=True)
        # 转换为pdf文件
        graph = pydotplus.graph_from_dot_data(dot_data) 
        graph.write_png('./result/图片/{0}_{1}_tree{2}.png'.format(pklpath,tiaocan_id,idx+1)) 

if __name__ == '__main__':
    #配置参数
    ###配置Word文档
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    ###配置字体
    mpl.rcParams['font.sans-serif'] = ['STZhongsong']    # 指定默认字体：解决plot不能显示中文问题
    mpl.rcParams['axes.unicode_minus'] = False   
    ###添加保存路径
    mkdir("./result/图片/")
    mkdir("./result/Word/")
    mkdir("./result/pkl/"+tiaocanmoshi[0])
    mkdir("./result/pkl/"+tiaocanmoshi[1])

    #加载数据
        #配置参数里面自动化加载

    #训练模型  
 
    itrname = "RF"
    pklpath=filepath
    if xunlian_flag == 0 :
        x_train,x_test,y_train,y_test,data_train,data_test,feature_name,max_features=data_divide(filepath,xunlian_flag)
        
        parameter = tiaocan(tiaocan_id) ###0简单调参局部最优  1网格搜索全局最优
        model_create(parameter['n_estimators'],parameter['max_depth'],parameter['max_features'],parameter['min_samples_split'],
                    parameter['min_samples_leaf'],parameter['criterion'])
        use_model(itrname,xunlian_flag)
    else :
        x_test,y_test,data_test,feature_name,max_features=data_divide(filepath,xunlian_flag)
        use_model(itrname,xunlian_flag)

    ###可视化
    shifou_printpng_tree=1        ###是否输出决策树0是 1否
    if shifou_printpng_tree== 0:
        filepath_temp = './data/pkl/_'+pklpath+'.pkl'
        itr=load_model(filepath_temp)
        keshihua(itr)
